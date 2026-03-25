#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标文件自动化生成引擎
支持公路、水利、市政房建、其他（监理/咨询）等工程类型
使用方式：
  python generate_bid.py --config project.json --output ./output
"""

import os
import sys
import json
import shutil
import argparse
import zipfile
import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] 缺少 python-docx 依赖，请运行: pip install python-docx")
    sys.exit(1)


# ─────────────────────────────────────────────
#  配置常量
# ─────────────────────────────────────────────
TEMPLATE_MAP = {
    "公路":     "templates/公路_模板.doc",
    "水利":     "templates/水利_模板.doc",
    "市政房建": "templates/市政房建_模板.doc",
    "其他":     "templates/其他_模板.doc",
}

# 公告模板（相对于本脚本目录）
NOTICE_TEMPLATE = "templates/招标公告_模板.docx"

# 可替换的占位符（模板中使用 {{KEY}} 形式）
PLACEHOLDER_FIELDS = [
    "PROJECT_NAME",       # 项目名称
    "PROJECT_CODE",       # 项目编号
    "OWNER_NAME",         # 招标人名称
    "OWNER_ADDRESS",      # 招标人地址
    "OWNER_CONTACT",      # 招标人联系方式
    "AGENT_NAME",         # 招标代理机构
    "AGENT_ADDRESS",      # 代理机构地址
    "AGENT_CONTACT",      # 代理机构联系方式
    "BID_SECTION",        # 标段名称
    "BID_AMOUNT",         # 估算金额
    "QUALIFY_LEVEL",      # 资质要求
    "DEADLINE_DATE",      # 投标截止日期
    "OPEN_DATE",          # 开标日期
    "OPEN_LOCATION",      # 开标地点
    "SELL_START_DATE",    # 招标文件发售起始日期
    "SELL_END_DATE",      # 招标文件发售截止日期
    "BOND_AMOUNT",        # 投标保证金金额
    "CONTRACT_PERIOD",    # 合同工期（天）
    "QUALITY_LEVEL",      # 质量标准
    "LOCATION",           # 工程地点
    "YEAR",               # 年份（自动填充）
    "MONTH",              # 月份（自动填充）
    "DAY",                # 日期（自动填充）
]


# ─────────────────────────────────────────────
#  工具函数
# ─────────────────────────────────────────────

def load_config(path: str) -> dict:
    """加载项目配置 JSON"""
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def _replace_in_paragraph(para, replacements: dict):
    """在段落中替换占位符（保持原有格式）"""
    # 收集段落完整文本
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    for key, value in replacements.items():
        placeholder = "{{" + key + "}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))

    # 将替换后的文本写回第一个 run，清除其余 run
    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ""


def replace_placeholders(doc: "Document", replacements: dict):
    """遍历文档所有段落与表格，替换占位符"""
    # 正文段落
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # 表格内段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)


def build_replacements(project: dict) -> dict:
    """构建占位符替换字典，自动填充日期字段"""
    now = datetime.now()
    replacements = {k: v for k, v in project.items()}
    replacements.setdefault("YEAR", str(now.year))
    replacements.setdefault("MONTH", str(now.month).zfill(2))
    replacements.setdefault("DAY", str(now.day).zfill(2))
    return replacements


def copy_template(src: str, dst: str):
    """复制模板文件到目标路径"""
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    shutil.copy2(src, dst)


def doc_to_docx(doc_path: str, docx_path: str) -> bool:
    """
    将 .doc 文件转为 .docx（需要 LibreOffice 或 soffice）
    返回 True 表示成功，False 表示跳过（直接使用 .doc）
    """
    import subprocess
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx",
             "--outdir", str(Path(docx_path).parent), doc_path],
            capture_output=True, timeout=60
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


# ─────────────────────────────────────────────
#  核心生成逻辑
# ─────────────────────────────────────────────

class BidDocGenerator:
    def __init__(self, config: dict, output_dir: str, templates_dir: str = "templates"):
        self.config = config
        self.output_dir = Path(output_dir)
        self.templates_dir = Path(templates_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _resolve_template(self, project_type: str, template_key: str = "main") -> Path:
        """解析模板文件路径"""
        # 优先使用配置中指定的模板
        custom = self.config.get("templates", {}).get(project_type)
        if custom and Path(custom).exists():
            return Path(custom)

        # 使用默认模板目录
        for ext in (".docx", ".doc"):
            candidate = self.templates_dir / f"{project_type}_模板{ext}"
            if candidate.exists():
                return candidate

        raise FileNotFoundError(
            f"找不到 [{project_type}] 类型的模板文件，"
            f"请在 templates/ 目录下放置 {project_type}_模板.docx"
        )

    def _resolve_notice_template(self) -> Path:
        """解析招标公告模板路径"""
        custom = self.config.get("notice_template")
        if custom and Path(custom).exists():
            return Path(custom)
        candidate = self.templates_dir / "招标公告_模板.docx"
        if candidate.exists():
            return candidate
        raise FileNotFoundError("找不到招标公告模板，请在 templates/ 目录下放置 招标公告_模板.docx")

    def generate_project(self, project: dict) -> dict:
        """
        为单个项目生成招标文件 + 招标公告
        返回生成的文件路径字典
        """
        name = project.get("PROJECT_NAME", "未命名项目")
        ptype = project.get("TYPE", "市政房建")
        idx = project.get("INDEX", "")
        prefix = f"{idx}." if idx else ""

        replacements = build_replacements(project)
        results = {}

        # ── 生成主体招标文件 ──
        try:
            tpl_path = self._resolve_template(ptype)
            out_name = f"{prefix}{name}.docx"
            out_path = self.output_dir / ptype / out_name
            out_path.parent.mkdir(parents=True, exist_ok=True)

            if tpl_path.suffix == ".doc":
                # 先复制，再尝试转换
                tmp_doc = out_path.with_suffix(".doc")
                shutil.copy2(tpl_path, tmp_doc)
                if doc_to_docx(str(tmp_doc), str(out_path)):
                    tmp_doc.unlink(missing_ok=True)
                    doc = Document(str(out_path))
                else:
                    # 无法转换，保持 .doc 格式
                    results["main"] = str(tmp_doc)
                    print(f"  [WARN] LibreOffice 未安装，已复制原始 .doc: {tmp_doc.name}")
                    doc = None
            else:
                doc = Document(str(tpl_path))

            if doc is not None:
                replace_placeholders(doc, replacements)
                doc.save(str(out_path))
                results["main"] = str(out_path)
                print(f"  [OK] 主体文件: {out_path.name}")

        except FileNotFoundError as e:
            print(f"  [SKIP] {e}")

        # ── 生成招标公告 ──
        try:
            notice_tpl = self._resolve_notice_template()
            notice_name = f"{prefix}{name}招标公告.docx"
            notice_path = self.output_dir / ptype / notice_name

            doc = Document(str(notice_tpl))
            replace_placeholders(doc, replacements)
            doc.save(str(notice_path))
            results["notice"] = str(notice_path)
            print(f"  [OK] 招标公告: {notice_path.name}")

        except FileNotFoundError as e:
            print(f"  [SKIP] {e}")

        return results

    def run(self):
        """批量生成所有项目文件"""
        projects = self.config.get("projects", [])
        if not projects:
            print("[WARN] 配置中没有找到任何项目（projects 字段为空）")
            return

        print(f"\n{'='*60}")
        print(f"  招标文件自动化生成  共 {len(projects)} 个项目")
        print(f"  输出目录: {self.output_dir.resolve()}")
        print(f"{'='*60}\n")

        success, failed = 0, 0
        for i, proj in enumerate(projects, 1):
            name = proj.get("PROJECT_NAME", f"项目{i}")
            print(f"[{i}/{len(projects)}] {name}")
            try:
                self.generate_project(proj)
                success += 1
            except Exception as e:
                print(f"  [ERROR] 生成失败: {e}")
                failed += 1

        print(f"\n{'='*60}")
        print(f"  完成！成功: {success}  失败: {failed}")
        print(f"{'='*60}\n")


# ─────────────────────────────────────────────
#  命令行入口
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="招标文件自动化生成工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_bid.py --config projects.json --output ./output
  python generate_bid.py --config projects.json --output ./output --templates ./templates
        """
    )
    parser.add_argument("--config",    required=True, help="项目配置 JSON 文件路径")
    parser.add_argument("--output",    default="./output", help="输出目录（默认 ./output）")
    parser.add_argument("--templates", default="./templates", help="模板目录（默认 ./templates）")
    args = parser.parse_args()

    if not os.path.exists(args.config):
        print(f"[ERROR] 配置文件不存在: {args.config}")
        sys.exit(1)

    config = load_config(args.config)
    gen = BidDocGenerator(config, args.output, args.templates)
    gen.run()


if __name__ == "__main__":
    main()
